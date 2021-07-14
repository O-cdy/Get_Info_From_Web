import os, docx, time, requests, urllib
from bs4 import BeautifulSoup
from selenium import webdriver
from selenium.webdriver.common.keys import Keys  
from selenium.webdriver.common.action_chains import ActionChains
from docx.shared import Pt, Inches #磅数,英尺
from docx.oxml.ns import qn #中文格式

#访问链接
url = "https://www.ximalaya.com/renwenjp/29030058/"
driver = webdriver.Firefox()
#driver = webdriver.Chrome()
driver.maximize_window()
driver.get(url)
#print(driver.page_source)

#登录按钮
elem_login = driver.find_element_by_xpath("/html/body/div[1]/header/div/div/div[2]/div/div/img")
print(elem_login)

#获取当前窗口句柄
#now_handle = driver.current_window_handle
#print(now_handle)
 
#获取当前登录按钮信息
print(elem_login.text)                 

driver.execute_script("arguments[0].click();", elem_login)

time.sleep(10)

#选择弹出的对话框
 
#获取用户名和密码(VIP account)
elem_name = driver.find_element_by_id("accountName")
elem_name.send_keys("XXXXXXXXXXX")
elem_pwd = driver.find_element_by_id("accountPWD")
elem_pwd.send_keys("XXXXXXXXXXX")

elem_pwd.send_keys(Keys.RETURN)
time.sleep(20)

#暂时未实现自动验证人机的功能，需手动验证非人机操作

#获得cookie信息
#cookiesInfo = driver.get_cookies()
#print(cookiesInfo)

# 各按钮css路径
page_path = 'html body div#award main.main-content div.album-detail div.clearfix div.detail.layout-main div#anchor_sound_list.sound-list-wrapper._is div.sound-list._is div.pagination._is nav.pagination.WJ_ ul.pagination-page.WJ_ li.page-item.WJ_ a.page-link.WJ_'
list_path = 'html body div#award main.main-content div.album-detail div.clearfix div.detail.layout-main div#anchor_sound_list.sound-list-wrapper._is div.sound-list._is ul li.lF_ div.text.lF_ a span.title.lF_'
note_path = 'html body div#award main.main-content div.sound-detail div.clearfix div.detail.layout-main div.sound-intro.eX_ article.intro.eX_ p'
image_path = 'html body div#award main.main-content div.sound-detail div.clearfix div.detail.layout-main div.sound-intro.eX_ article.intro.eX_ p img'

doc = BeautifulSoup(driver.page_source, 'html.parser')
page_selector = doc.select(page_path)
print(page_selector)


#创建内存中的word文档对象
file=docx.Document()
file.styles['Normal'].font.name = u'宋体'
file.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')


time.sleep(10)
x=0

for page_num in range(0,5): 
	page_result = page_selector[page_num]
	#print(page_result)
	next_url = page_result['href']
#找到对应页面按钮位置
	page_button = driver.find_element_by_link_text(page_result.text)
	print(page_button)
	print(page_button.text)
#	now_handle = driver.current_window_handle
	#点击按钮
	driver.execute_script("arguments[0].click();", page_button)
	time.sleep(5)
	doc1 = BeautifulSoup(driver.page_source, 'html.parser')
	next_result = doc1.select(list_path)
	while(next_result == []):
		next_result = doc1.select(list_path)
#	print(next_result)
	for item in next_result:
		print(item.text)
		item_detail = driver.find_element_by_link_text(item.text)
		driver.execute_script("arguments[0].click();", item_detail)
		time.sleep(5)
		driver.refresh()
		time.sleep(5)
		doc2 = BeautifulSoup(driver.page_source, 'html.parser')
		note_result = doc2.select(note_path)
		cookiesInfo = driver.get_cookies()
		#print(cookiesInfo)
		#print(note_result)
		file.add_paragraph(item.text)
		image_result = []
		image_result = doc2.select(image_path)
#		print(image_result)
		path='D:\\microsoft\\pic'
		while(image_result != []):
			for image in image_result:
				print(image['src'])
				if not os.path.isdir(path):
					os.makedirs(path)
				download_url = image['src']
				urllib.request.urlretrieve(download_url,'D:\\microsoft\\pic\\'+str(x)+'.jpg')
				width = Inches(6)
				height = Inches(4)
				file.add_picture('D:\\microsoft\\pic\\'+str(x)+'.jpg',width = Inches(6))
				x=x+1
				print(x)
			break
		for note in note_result:
			print(note.text)
#			#写入若干段落
			file.add_paragraph(note.text)
			#file.add_paragraph().space_after = Pt(10) #段后空间
		#分页
		file.add_page_break()
		print('')
		#找到回到上一页的按钮
		z = driver.find_element_by_link_text('人文社科经典必读80本 | 喜马讲书')
		driver.execute_script("arguments[0].click();", z)
		time.sleep(5)
		#重新定位到对应的页数
		page_button = driver.find_element_by_link_text(page_result.text)
		driver.execute_script("arguments[0].click();", page_button)
		driver.refresh()
		time.sleep(5)

#文件存储
file.save("D:\microsoft\笔记_final.docx")
driver.delete_all_cookies()
driver.quit()
		


##################
#
#  问题： 
#    1. C盘爆满参考：https://blog.csdn.net/weixin_30773135/article/details/96187876
#	 2. 有时网页加载失败，报错500 --> 解决方法： time.sleep(5) 配合 driver.refresh()，确保网页信息完全加载
# 	 3. 登录时，机器人验证的环节暂未实现自动化
#	 4. 为了正确取到结果，避免网页中，就从页面提取信息的情况，使程序sleep(5)，让页面元素加载完全
#	 5. elem_login.click() 时常失败，解决方法： driver.execute_script("arguments[0].click();", elem_login)
#	 6. 整体数据提取速度有待提升
#
#
#  笔记：
#	 1. 获得cookie信息：cookiesInfo = driver.get_cookies()
#	 2. word文档对象
#	 	创建 : file=docx.Document()
#	 	读取 ：file=docx.Document(path)
#
#
#####################


